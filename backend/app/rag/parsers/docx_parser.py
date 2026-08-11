"""DOCX parser for paragraphs, headings, tables and embedded images."""

from __future__ import annotations

import json
import os
import zipfile
from pathlib import Path

from ..blocks import DocumentBlock
from ..ocr import create_ocr_engine
from ..vision import VisionAnalyzer
from .base import BaseParser
from .formula import append_formulas, extract_native_formulas
from .media import analyze_media


class DocxParser(BaseParser):
    source_type = "docx"

    def __init__(
        self,
        document_id: str,
        ocr_engine=None,
        vision_analyzer: VisionAnalyzer | None = None,
        formula_recognizer=None,
        work_dir: str | None = None,
        original_dir: str | None = None,
    ):
        super().__init__(document_id)
        # original_dir 兼容统一接口（worker 对所有解析器传该参数）；DOCX 原文件在管理目录，无需复制
        self.ocr_engine = ocr_engine or create_ocr_engine()
        self.vision_analyzer = vision_analyzer
        self.formula_recognizer = formula_recognizer
        self.work_dir = Path(work_dir or os.getenv("RAG_WORK_DIR", "data/.work"))

    def parse(self, path: str | Path) -> list[DocumentBlock]:
        from docx import Document

        document = Document(str(path))
        blocks: list[DocumentBlock] = []
        headings: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            formulas = extract_native_formulas(paragraph._p)
            text = append_formulas(text, formulas)
            style_name = (paragraph.style.name or "") if paragraph.style else ""
            if style_name.lower().startswith("heading"):
                try:
                    level = int(style_name.rsplit(" ", 1)[-1])
                except ValueError:
                    level = 1
                headings = headings[: max(level - 1, 0)]
                if text:
                    headings.append(text)
                    blocks.append(
                        self._block(
                            text,
                            content_type="formula" if formulas else "heading",
                            heading_path=headings,
                            metadata={"formulas_latex": formulas} if formulas else {},
                        )
                    )
                continue
            if text:
                blocks.append(
                    self._block(
                        text,
                        content_type="formula" if formulas else "text",
                        heading_path=headings,
                        metadata={"formulas_latex": formulas} if formulas else {},
                    )
                )

        for table_index, table in enumerate(document.tables):
            rows = []
            formula_rows = []
            for row in table.rows:
                row_values = []
                row_formulas = []
                for cell in row.cells:
                    formulas = extract_native_formulas(cell._tc)
                    row_values.append(append_formulas(cell.text, formulas))
                    row_formulas.append(formulas)
                rows.append(row_values)
                formula_rows.append(row_formulas)
            if not rows:
                continue
            table_text = self._table_to_text(rows)
            blocks.append(
                self._block(
                    table_text,
                    content_type="table",
                    heading_path=headings,
                    metadata={
                        "table": rows,
                        "table_index": table_index,
                        "formulas_latex": formula_rows,
                    },
                )
            )

        blocks.extend(self._parse_embedded_images(path))
        return blocks

    def _parse_embedded_images(self, path: str | Path) -> list[DocumentBlock]:
        image_dir = self.work_dir / self.document_id / "docx_images"
        blocks: list[DocumentBlock] = []
        with zipfile.ZipFile(path) as archive:
            members = [name for name in archive.namelist() if name.startswith("word/media/")]
            for index, member in enumerate(members):
                suffix = Path(member).suffix.lower() or ".png"
                image_path = image_dir / f"image_{index}{suffix}"
                image_path.parent.mkdir(parents=True, exist_ok=True)
                image_path.write_bytes(archive.read(member))
                text, metadata, confidence, content_type = analyze_media(
                    image_path,
                    self.ocr_engine,
                    self.vision_analyzer,
                    self.formula_recognizer,
                )
                metadata.update({"embedded_member": member, "asset_kind": "docx_image"})
                blocks.append(
                    self._block(
                        text,
                        content_type=content_type,
                        image_path=str(image_path),
                        confidence=confidence,
                        metadata=metadata,
                    )
                )
        return blocks

    @staticmethod
    def _table_to_text(rows: list[list[str]]) -> str:
        if not rows:
            return ""
        return "Word 表格：\n" + json.dumps(rows, ensure_ascii=False)
